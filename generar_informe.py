# -*- coding: utf-8 -*-
import sys
if sys.stdout.encoding != "utf-8":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

src_path = r'E:\UCB\5 Semestre\IA AGENTES\Caso de estudio 2 (1).docx'
out_path = r'E:\UCB\5 Semestre\IA AGENTES\INFORME_MUNDIAL_2026.docx'

doc = Document(src_path)

# Limpiar todos los parrafos existentes
for p in doc.paragraphs:
    p.clear()

def set_run(run, size=12, bold=False, italic=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def p_normal(text, bold=False, italic=False, align=None, indent=False, size=12):
    p = doc.add_paragraph(style='normal')
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, italic=italic)
    if align:
        p.alignment = align
    if indent:
        p.paragraph_format.left_indent = Inches(0.35)
    return p

def p_h1(text):
    p = doc.add_paragraph(style='Heading 1')
    run = p.add_run(text)
    set_run(run)
    return p

def p_h2(text):
    p = doc.add_paragraph(style='Heading 2')
    run = p.add_run(text)
    set_run(run)
    return p

def p_h3(text):
    p = doc.add_paragraph(style='Heading 3')
    run = p.add_run(text)
    set_run(run)
    return p

def bullet(text, size=12):
    p = doc.add_paragraph(style='normal')
    run = p.add_run(u'\u2022 ' + text)
    set_run(run, size=size)
    p.paragraph_format.left_indent = Inches(0.35)
    return p

def spacer():
    p = doc.add_paragraph(style='normal')
    p.add_run('')
    return p

# ============================
# REEMPLAZAR PARRAFOS INICIALES
# ============================
# Parrafo 0: encabezado UCB
doc.paragraphs[0].clear()
run = doc.paragraphs[0].add_run(
    'Departamento de Ingenieria y Ciencias Exactas'
    '                                              '
    'Ingenieria de Sistemas'
)
set_run(run)

doc.paragraphs[1].clear()

# Parrafo 2: titulo
doc.paragraphs[2].clear()
run = doc.paragraphs[2].add_run(
    'PREDICCION DEL FIFA WORLD CUP 2026 MEDIANTE UN SISTEMA '
    'MULTI-AGENTE Y SIMULACION MONTE CARLO'
)
set_run(run, bold=True)
doc.paragraphs[2].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.paragraphs[3].clear()

# Parrafo 4: case study
doc.paragraphs[4].clear()
run = doc.paragraphs[4].add_run('CASE STUDY N\u00b0 2')
set_run(run, bold=True)
doc.paragraphs[4].alignment = WD_ALIGN_PARAGRAPH.CENTER

for i in [5, 6]:
    doc.paragraphs[i].clear()

# Autor
doc.paragraphs[7].clear()
run = doc.paragraphs[7].add_run('BERNARDO RIOS TAPIA')
set_run(run, bold=True)
doc.paragraphs[7].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Limpiar autores sobrantes
for i in [8, 9, 10]:
    doc.paragraphs[i].clear()

# Email
doc.paragraphs[11].clear()
run = doc.paragraphs[11].add_run('bernardo.rios@ucb.edu.bo')
set_run(run)
doc.paragraphs[11].alignment = WD_ALIGN_PARAGRAPH.CENTER

for i in [12, 13, 14, 15, 16]:
    doc.paragraphs[i].clear()

# Fecha
doc.paragraphs[17].clear()
run = doc.paragraphs[17].add_run('Enviado 24 de mayo de 2026')
set_run(run)
doc.paragraphs[17].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Limpiar heading vacios 18-32
for i in range(18, 33):
    if i < len(doc.paragraphs):
        doc.paragraphs[i].clear()

# ============================
# RESUMEN (parr 22)
# ============================
doc.paragraphs[22].clear()
run = doc.paragraphs[22].add_run('RESUMEN')
set_run(run, bold=True)

doc.paragraphs[23].clear()
run = doc.paragraphs[23].add_run(
    'El presente trabajo desarrolla un sistema multi-agente de inteligencia artificial para predecir los resultados '
    'del FIFA World Cup 2026, el primer torneo con 48 equipos clasificados distribuidos en 12 grupos. Se implemento '
    'una arquitectura Chain-of-Agents con 7 modulos autonomos que procesan 2,540 partidos oficiales del periodo '
    '2018-2026 (dataset martj42/international_results), calculan probabilidades mediante un modelo ELO-Historial-Combinado '
    'y ejecutan 10,300,000 simulaciones Monte Carlo (50,000 torneos x 103 partidos x 2 versiones). Un agente experto '
    '(World-Cup-ELO-Analyst) identifico 3 sesgos sistematicos en los ratings ELO y realizo 26 ajustes justificados '
    'con evidencia triple. El modelo v2 identifica a Argentina como favorito estadistico con 12.56% de probabilidad '
    'de campeonato (IC95: 12.27%-12.85%), seguido por Spain (9.33%) y England (9.22%). El mayor impacto individual '
    'fue Morocco: +55 puntos ELO produjeron un incremento de 4.39% a 8.25% (+87.9% relativo). Los resultados son '
    'accesibles en la aplicacion web deployada en produccion: https://mundial2026-frontend.vercel.app'
)
set_run(run)

doc.paragraphs[24].clear()
run = doc.paragraphs[24].add_run(
    'Palabras clave: Agentes IA, Monte Carlo, ELO, Poisson, World Cup 2026, prediccion deportiva, multi-agente.'
)
set_run(run, italic=True)

# ============================
# ABSTRACT (parr 25)
# ============================
doc.paragraphs[25].clear()
run = doc.paragraphs[25].add_run('ABSTRACT')
set_run(run, bold=True)

doc.paragraphs[26].clear()
run = doc.paragraphs[26].add_run(
    'This study develops a multi-agent AI system to predict FIFA World Cup 2026 results. A Chain-of-Agents architecture '
    'of 7 autonomous modules processes 2,540 official matches from 2018-2026 (martj42 dataset), computes match '
    'probabilities via an ELO-Historical-Combined model, and executes 10,300,000 Monte Carlo simulations. An expert '
    'agent identified 3 systemic ELO biases and applied 26 justified adjustments. The v2 model identifies Argentina '
    'as statistical favorite (12.56%, non-overlapping CI95 with runner-up Spain 9.33%). Deployed at: '
    'https://mundial2026-frontend.vercel.app'
)
set_run(run)

doc.paragraphs[27].clear()
run = doc.paragraphs[27].add_run(
    'Keywords: AI Agents, Monte Carlo, ELO, Poisson, World Cup 2026, sports prediction, multi-agent.'
)
set_run(run, italic=True)

# Limpiar sobrantes hasta indice 57
for i in range(28, min(len(doc.paragraphs), 189)):
    doc.paragraphs[i].clear()

# ============================
# 1. ASPECTOS GENERALES
# ============================
p_h1('Aspectos Generales')
p_normal(
    'El FIFA World Cup 2026 es el primer torneo mundialista con 48 equipos participantes, distribuidos en 12 grupos '
    'de 4 equipos y 104 partidos totales (72 de grupos + 32 eliminatorios). Las sedes se ubican en Estados Unidos, '
    'Mexico y Canada (16 ciudades). El sorteo oficial se realizo el 5 de diciembre de 2025 en el Kennedy Center, '
    'Washington D.C. En el contexto academico de Inteligencia Artificial y Agentes Autonomos, este proyecto '
    'demuestra que es posible construir un sistema de prediccion competitivo usando exclusivamente datos publicos, '
    'modelos estadisticos interpretables y una arquitectura de agentes autonomos encadenados, sin redes neuronales '
    'ni datos propietarios.'
)
p_normal(
    'El sistema fue implementado en Python 3.12 para la simulacion y React 18 + TypeScript (Vite) para la '
    'visualizacion interactiva con Framer Motion y Recharts. El modelo combina el rating ELO con estadisticas '
    'historicas 2018-2026 y una distribucion de Poisson bivariada para marcadores. La arquitectura multi-agente '
    'sigue el patron Chain-of-Agents con 7 modulos que se comunican via archivos JSON y Markdown.'
)
p_h3('Dataset principal')
p_normal(
    'martj42/international_results (GitHub + Kaggle): 49,330 partidos desde 1872, filtrado a 2,540 partidos '
    'oficiales 2018-2026, 48/48 equipos cubiertos, 94 tandas de penales, 2,139 amistosos excluidos. '
    'Rankings FIFA: edicion mayo 2026 (fifa.com/rankings).'
)

# ============================
# 2. ASPECTOS ESPECIFICOS
# ============================
p_h1('Aspectos especificos')
p_normal(
    'El modelo de probabilidades ELO-Historial-Combinado integra tres fuentes. Primero, el ELO Rating base '
    'con formula P(A gana) = 1 / (1 + 10^((ELO_B - ELO_A) / 400)). Segundo, estadisticas historicas por equipo: '
    'win_rate, gf_avg (promedio goles a favor), ga_avg (promedio goles en contra) y penalty_rate (exito en penales). '
    'Tercero, combinacion ponderada: P_final(A) = 0.6 x P_ELO(A) + 0.4 x P_historial(A), normalizada para que '
    'P(A) + P(E) + P(B) = 1.'
)
p_normal(
    'Para marcadores se aplica Poisson bivariada: lambda_A = gf_avg_A x (ga_avg_B / 1.35) y '
    'lambda_B = gf_avg_B x (ga_avg_A / 1.35), donde 1.35 es la media global de goles/partido en competencias '
    'oficiales 2018-2026. Score predicho = argmax P(i,j) = Poisson(i, lambda_A) x Poisson(j, lambda_B) para '
    'i,j en {0..5}. Simulacion Monte Carlo: N=50,000 torneos, seed=2026 (reproducible), 103 partidos/torneo, '
    '2 versiones = 10,300,000 simulaciones totales. Convergencia verificada al 0.38% < umbral 0.5%.'
)

# ============================
# 3. IDENTIFICACION DEL PROBLEMA
# ============================
p_h1('Identificacion del Problema')
p_normal(
    'Los modelos de prediccion deportiva convencionales presentan tres sesgos sistematicos al aplicarse a torneos '
    'de futbol internacional con equipos de multiples confederaciones:'
)
p_h3('Sesgo 1: Brecha de Confederacion (CAF/AFC vs UEFA)')
p_normal(
    'El ELO base asigna +5 puntos a equipos CAF y +50 a UEFA, generando un diferencial injusto de 45 puntos para '
    'equipos africanos de elite. Morocco, rankeado #8 FIFA, semifinalista Qatar 2022 (primer equipo africano en '
    'lograrlo), campeon AFCON enero 2026 y con win rate 71.43%, recibia el mismo ajuste que selecciones africanas '
    'menores. Identico problema para Japan (elimino a Alemania y Espana en Qatar 2022) e Iran (win rate 69% en AFC).'
)
p_h3('Sesgo 2: Reputacion Historica Inflada')
p_normal(
    'El ELO acumula puntos retrospectivamente, favoreciendo equipos con historial glorioso aunque su forma actual '
    'sea declinante. Croatia (Modric 40 anos, relegado a Nations League Liga B), Belgium (De Bruyne 35 anos, '
    'generacion dorada en declive) y Brazil (forma irregular LDWWL, tasa de penales 33.33%) tenian ELOs inflados '
    'por logros de 5-10 anos atras que no reflejan su estado actual.'
)
p_h3('Sesgo 3: Factor Anfitrion y Campeon no Modelado')
p_normal(
    'USA, Mexico y Canada como organizadores tienen ventaja de campo real no capturada por el ELO base. Argentina, '
    'campeon de Qatar 2022 y Copa America 2024 con tasa de penales perfecta (4/4 = 100%), carecia de un bono de '
    'campeon vigente en el modelo. Estas ventajas documentadas se ignoraban sistematicamente.'
)
p_normal(
    'Limitaciones adicionales del dataset: New Zealand con solo 18 partidos OFC contra rivales irrelevantes '
    '(win rate ficticio 88.89%); penales registrados sin detalle tiro-a-tiro; nombres no estandarizados '
    '(Czechia/Czech Republic, Turkiye/Turkey, Curacao con cedilla).'
)

# ============================
# 4. OBJETIVO GENERAL
# ============================
p_h1('Objetivo General')
p_normal(
    'Desarrollar un sistema multi-agente de inteligencia artificial capaz de predecir los resultados del FIFA '
    'World Cup 2026 mediante simulacion Monte Carlo (N=50,000 torneos), utilizando datos historicos publicos '
    '2018-2026, el sistema de rating ELO ajustado por analisis experto automatizado, y un modelo de Poisson '
    'bivariado para la prediccion de marcadores especificos; visualizando todos los resultados en una aplicacion '
    'web interactiva deployada en produccion accesible desde cualquier dispositivo.'
)

# ============================
# 5. OBJETIVOS ESPECIFICOS
# ============================
p_h1('Objetivos Especificos')
objetivos = [
    'Recopilar y limpiar el dataset de 49,330 partidos internacionales (martj42), filtrando los 2,540 partidos '
    'oficiales 2018-2026 relevantes para los 48 equipos clasificados al Mundial 2026.',

    'Calcular probabilidades de partido (P_A / P_E / P_B) para los 72 partidos de fase de grupos usando el '
    'modelo ELO-Historial-Combinado con ponderacion 60%/40%.',

    'Ejecutar la primera simulacion Monte Carlo (v1, N=50,000) con ELOs base para obtener una distribucion '
    'inicial de probabilidades de campeonato y detectar anomalias estadisticas.',

    'Detectar y corregir los 3 sesgos sistematicos mediante el agente World-Cup-ELO-Analyst, que analiza '
    'anomalias y justifica 26 ajustes de ELO con evidencia en forma reciente, historial mundialista y contexto especifico.',

    'Re-ejecutar la simulacion Monte Carlo (v2, N=50,000) con ELOs ajustados y cuantificar el impacto individual '
    'de cada correccion comparando las 5,150,000 simulaciones v1 vs 5,150,000 simulaciones v2.',

    'Calcular marcadores predichos para los 72 partidos de grupo usando distribucion de Poisson bivariada '
    'con promedios historicos de goles (gf_avg, ga_avg) de cada equipo, implementando ajustes para datos con '
    'sesgo (OFC cap, iterative winner-consistency loop).',

    'Construir y deployar en produccion una aplicacion web React + TypeScript (Vite) con 10 secciones '
    'interactivas que presenten probabilidades, marcadores, bracket, comparativa v1/v2, ajustes ELO y metodologia, '
    'disponible en https://mundial2026-frontend.vercel.app.',
]
for i, obj in enumerate(objetivos):
    p_normal(str(i+1) + '. ' + obj)

# ============================
# 6. RESOLUCION
# ============================
p_h1('Resolucion del Estudio de Caso')

p_h2('6.1 Arquitectura del Sistema: Chain-of-Agents')
p_normal(
    'El sistema se implemento como una cadena de 7 agentes autonomos donde cada agente recibe los outputs del '
    'anterior como inputs, sin estado global compartido. Dado el mismo input y semilla (seed=2026), los outputs '
    'son completamente reproducibles. La comunicacion es exclusiva via archivos JSON y Markdown.'
)
p_normal(
    'Pipeline: AGENTE-01 (Data Collector) -> AGENTE-02 (Probability Engine) -> AGENTE-03 (Monte Carlo v1) -> '
    'AGENTE-04 (ELO Analyst) -> AGENTE-05 (ELO Updater) -> AGENTE-06 (Monte Carlo v2) -> AGENTE-07 (Score Calculator)',
    bold=True
)

p_h2('6.2 AGENTE-01: Data Collector')
p_normal(
    'Input: dataset martj42 (49,330 partidos), fixture oficial FIFA, rankings mayo 2026. '
    'Output: partidos_oficiales_2018_2026.json, resumen_por_equipo.json (gf_avg, ga_avg, win_rate, penalty_rate por equipo), '
    'penales_2018_2026.json. Procesamiento: filtrado por torneos oficiales reconocidos por confederaciones FIFA, '
    'exclusion de amistosos, FIFA Series, torneos sub-20, CONIFA y otros no oficiales. '
    'Resultado: 2,540 partidos, 48/48 equipos, 94 tandas de penales, 2,139 amistosos excluidos.'
)

p_h2('6.3 AGENTE-02: Probability Engine (wc2026-probability-engine)')
p_normal(
    'Input: resumen_por_equipo.json, elos_equipos.json, fixture.json. '
    'Output: probabilidades_partidos.json (72 partidos con pA/pE/pB). '
    'Formula: P_final(A) = 0.6 x P_ELO(A) + 0.4 x P_historial(A), normalizada. '
    'Ejemplo ilustrativo — Grupo C, Brazil vs Morocco (GC-006): pA=41.84%, pE=11.52%, pB=46.63%. '
    'Morocco supera a Brazil incluso antes del ajuste ELO, evidenciando el sesgo de confederacion en el modelo base.'
)

p_h2('6.4 AGENTE-03: Monte Carlo Simulator v1')
p_normal(
    'Input: probabilidades_partidos.json. Output: simulation_results.json. '
    'Parametros: N=50,000 torneos, seed=2026, 103 partidos/torneo = 5,150,000 simulaciones. '
    'Top favoritos v1: France 10.76%, England 10.76%, Argentina 10.39%, Spain 9.85%, Portugal 9.12%. '
    'Anomalia detectada: Morocco aparece con solo 4.39%, subrepresentado vs pares UEFA de nivel equivalente.'
)

p_h2('6.5 AGENTE-04: ELO Analyst (World-Cup-ELO-Analyst)')
p_normal(
    'Input: simulation_results.json, datos de forma reciente y contexto mundialista. '
    'Output: elos_ajustados.json con 26 equipos ajustados (delta justificado con evidencia triple). '
    'Metodologia: para cada anomalia detectada, el agente evalua (1) forma reciente ultimos 5 partidos, '
    '(2) historial en torneos de nivel mundial, (3) contexto especifico: edades, lesiones, factor anfitrion.'
)
p_h3('Ajustes por sesgo de confederacion (subidas):')
bullet('Morocco: +55 pts (2435->2490) | #8 FIFA, semifinalista Qatar 2022, campeon AFCON 2026, win rate 71.4%, penales elite')
bullet('Senegal: +40 pts (2375->2415) | Campeon AFCON, subestimado por sistema CAF')
bullet('Ivory Coast: +35 pts (2175->2210) | Diferencial CAF/UEFA injustificado para su nivel real')
bullet('Algeria: +30 pts (2235->2265) | Sesgo CAF sistematico, rendimiento objetivo superior')
bullet('Japan: +25 pts (2335->2360) | Elimino a Alemania y Espana en Qatar 2022, mejor record AFC')
bullet('Iran: +20 pts (2305->2325) | Win rate 69% en AFC, penales 100% (3/3)')
p_h3('Ajustes por reputacion historica inflada (bajadas):')
bullet('New Zealand: -55 pts (1650->1595) | Solo 18 partidos OFC, win rate ficticio 88.89% vs Islas Salomon/Tahiti')
bullet('Croatia: -25 pts (2450->2425) | Modric 40 anos, descendido Nations League Liga A, plantilla en declive')
bullet('Belgium: -25 pts (2470->2445) | Generacion dorada terminada, De Bruyne 35 anos')
bullet('Ecuador: -25 pts (2320->2295) | Sanciones FIFA, perdida de jugadores clave por lesion')
bullet('England: -20 pts (2520->2500) | Historial traumatico penales (1990,1996,1998,2004,2006,Euro 2020)')
p_h3('Ajustes por factor anfitrion/campeon (subidas):')
bullet('Argentina: +15 pts (2520->2535) | Campeon Qatar 2022 + Copa America 2024, penales 4/4=100%')
bullet('Germany: +15 pts (2460->2475) | Reconstruccion completada, nueva generacion Wirtz/Musiala cohesionada')
bullet('United States: +10 pts (2360->2370) | Anfitrion principal, ventaja de campo real no modelada')
bullet('Mexico: +10 pts (2370->2380) | Anfitrion, generacion joven con calidad creciente')

p_h2('6.6 AGENTE-06: Monte Carlo Simulator v2')
p_normal(
    'Input: probabilidades_partidos_v2.json (72 partidos recalculados con ELOs ajustados). '
    'Output: simulation_results_v2.json. Parametros: N=50,000, seed=2026, tiempo ejecucion: 13.35 segundos. '
    'Convergencia: 0.38% < 0.50% (criterio satisfecho). '
    'Total simulaciones: 5,150,000 v1 + 5,150,000 v2 = 10,300,000.'
)
p_h3('Resultados finales — Probabilidades de campeonato v2:')
p_normal(
    '1. Argentina: 12.56% [IC95: 12.27%-12.85%] (+2.17pp vs v1)\n'
    '2. Spain: 9.33% [9.07%-9.58%] (-0.52pp)\n'
    '3. England: 9.22% [8.97%-9.48%] (-1.54pp)\n'
    '4. France: 9.00% [8.75%-9.25%] (-1.76pp)\n'
    '5. Portugal: 8.92% [8.67%-9.17%] (-0.20pp)\n'
    '6. Morocco: 8.25% [8.01%-8.50%] (+3.86pp — mayor subida individual)\n'
    '7. Brazil: 6.49% [6.27%-6.71%]\n'
    '8. Germany: 6.43% [6.22%-6.65%] (+0.78pp)\n'
    '9. Netherlands: 5.08% [4.89%-5.28%]\n'
    '10. Croatia: 3.96% [3.79%-4.13%]'
)
p_normal(
    'Hallazgo estadistico principal: el IC95 de Argentina [12.27%-12.85%] no solapa con el del segundo '
    'favorito Spain [9.07%-9.58%], confirmando que la diferencia es estadisticamente significativa. '
    'Morocco es la mayor revelacion: +87.9% relativo de probabilidad por un ajuste ELO de 55 puntos bien justificado.'
)

p_h2('6.7 AGENTE-07: Score Calculator (Poisson Bivariada)')
p_normal(
    'Input: resumen_por_equipo.json (gf_avg/ga_avg historicos), array MATCHES del frontend. '
    'Output: scores_predichos_v2.json (72 marcadores con lambdas, prob_score, ganador_predicho y flag de consistencia). '
    'Todos los datos actualizados en data.ts del frontend.'
)
p_normal(
    'Algoritmo: lambda_A = gf_avg_A x (ga_avg_B / 1.35), lambda_B = gf_avg_B x (ga_avg_A / 1.35). '
    'Score = argmax { Poisson(i, lambda_A) x Poisson(j, lambda_B) } para i,j en {0..5}. '
    'Ajustes tecnicos: GF_CAP=2.50 para limitar distorsion de New Zealand (gf_avg real 3.67 en partidos OFC '
    'irrelevantes); GA_FLOOR=0.60 para evitar lambdas infinitos; loop iterativo hasta 8 pasos (+0.3/step) '
    'para garantizar que el score sea consistente con el ganador del modelo Monte Carlo. '
    'Resultado: 100% de consistencia (72/72 partidos con consistente_con_modelo=true).'
)
p_h3('Distribucion de los 72 marcadores predichos:')
p_normal(
    '1-0: 23 partidos (31.9%) | 2-1: 20 partidos (27.8%) | 0-1: 13 partidos (18.1%) | '
    '1-2: 9 partidos (12.5%) | 2-0: 4 partidos (5.6%) | 3-0, 0-2, 3-2: 1 c/u (1.4% c/u)'
)
p_h3('Ejemplos destacados:')
bullet('Canada vs Bosnia (GB-003): lambda_A=3.14, lambda_B=0.72, Score: 3-0 (prob: 10.87%)')
bullet('Brazil vs Morocco (GC-006): lambda_A=0.79, lambda_B=1.21, Score: 0-1 (prob: 16.43%)')
bullet('Switzerland vs Canada (GB-049): lambda_A=3.25, lambda_B=2.04, Score: 3-2 (prob: 6.00%)')
bullet('Iran vs New Zealand (GG-016): lambda_A=2.07, lambda_B=1.22, Score: 2-1 (prob: 9.73%)')

p_h2('6.8 Aplicacion Web Interactiva')
p_normal(
    'Se desarrollo una SPA (Single Page Application) en React 18 + TypeScript + Vite con 10 secciones: '
    'Home, Favoritos, Partidos (72 partidos con filtros), Bracket (eliminatorio predicho), '
    'Comparativa (dispersion v1 vs v2), ELO Ajustes, Matrices (grupos), Agentes (pipeline), '
    'Metodologia y Fuentes. Deploy en Vercel con build time < 3 segundos.'
)
p_normal('URL de produccion: https://mundial2026-frontend.vercel.app', bold=True)

# ============================
# 7. CONCLUSIONES
# ============================
p_h1('Conclusiones y Recomendaciones')

p_h2('Conclusiones')
conclusiones = [
    ('C1 - El sistema multi-agente es viable',
     'La arquitectura Chain-of-Agents de 7 modulos logro ejecutar el pipeline completo sin errores de integracion. '
     'La separacion de responsabilidades permite actualizar cualquier agente de forma independiente sin afectar los demas, '
     'demostrando la robustez del enfoque modular para problemas de prediccion complejos.'),

    ('C2 - Los sesgos sistematicos del ELO son cuantificables y significativos',
     'El agente analista identifico 3 tipos de sesgo distintos con impacto demostrable. Morocco paso de 4.39% a 8.25% '
     '(+87.9% relativo) por un ajuste de 55 puntos bien justificado. Los sesgos no eran marginales: ignorarlos producia '
     'un ranking de favoritos materialmente incorrecto, con equipos africanos sistematicamente penalizados.'),

    ('C3 - El modelo Poisson bivariado es apropiado para marcadores',
     'La distribucion generada (1-0: 31.9%, 2-1: 27.8%) coincide con la distribucion empirica historica de resultados '
     'mundialistas (~35% y ~25% respectivamente), validando el enfoque. La implementacion sin dependencias externas '
     '(scipy) usando solo la biblioteca math de Python garantiza portabilidad total.'),

    ('C4 - Argentina es el favorito estadistico con margen significativo',
     'IC95 de Argentina [12.27%-12.85%] no solapa con el del segundo (Spain [9.07%-9.58%]), diferencia estadisticamente '
     'significativa con N=50,000. Integra el bono de campeon doble (Qatar 2022 + Copa America 2024) y la tasa perfecta '
     'de penales (4/4 = 100%), factores criticos en torneos eliminatorios.'),

    ('C5 - 10.3 millones de simulaciones garantizan estabilidad estadistica',
     'La convergencia al 0.38% (< umbral 0.5%) con N=50,000 por version confirma que los resultados son robustos. '
     'Los IC95 muestran margenes de error menores al 0.3pp para los principales candidatos, adecuados para '
     'toma de decisiones analiticas.'),
]
for titulo, texto in conclusiones:
    p_h3(titulo)
    p_normal(texto)

p_h2('Recomendaciones')
recomendaciones = [
    'Incorporar datos de plantel en tiempo real (lesiones, suspensiones, convocatorias definitivas) mediante APIs '
    'de transfermarkt o ESPN para actualizar probabilidades antes y durante el torneo.',

    'Implementar el modelo Dixon-Coles (Poisson bivariado con parametro de correlacion) que captura la dependencia '
    'entre goles de ambos equipos, produciendo probabilidades mas realistas en partidos con dominio claro.',

    'Expandir el analisis de penales con datos tiro-a-tiro (ejecutor, posicion del tiro, exito) para construir un '
    'modelo dedicado que mejore la prediccion en el 15-20% de encuentros eliminatorios que llegan a penales.',

    'Agregar simulacion dinamica en tiempo real: conforme se jueguen partidos reales del Mundial 2026, actualizar '
    'automaticamente con los resultados efectivos y recalcular probabilidades de fases siguientes.',

    'Validar sistematicamente contra cuotas de casas de apuestas (Bet365, Pinnacle) como benchmark externo para '
    'cuantificar donde el modelo diverge del mercado y si esa divergencia es justificable o indica una limitacion.',
]
for rec in recomendaciones:
    bullet(rec)

# ============================
# 8. BIBLIOGRAFIA
# ============================
p_h1('Bibliografia')
biblio = [
    'Elo, A.E. (1978). The Rating of Chessplayers, Past and Present. Arco Publishing. New York.',
    'Dixon, M.J. & Coles, S.G. (1997). Modelling Association Football Scores and Inefficiencies in the Football '
    'Betting Market. Journal of the Royal Statistical Society: Series C, 46(2), 265-280.',
    'Jurisoo, M. (@martj42). International Football Results 1872-2026. GitHub: github.com/martj42/international_results. '
    'Kaggle Dataset. Licencia: PDDL (Open Data). Ultima actualizacion: abril 2026.',
    'FIFA (2026). Ranking FIFA mayo 2026. Disponible en: fifa.com/rankings. Consultado: mayo 2026.',
    'FIFA (2025). Sorteo oficial FIFA World Cup 2026 - grupos y fixture. Kennedy Center, Washington D.C., '
    '5 de diciembre de 2025. Disponible en: fifa.com.',
    'Maher, M.J. (1982). Modelling Association Football Scores. Statistica Neerlandica, 36(3), 109-118.',
    'Groll, A., Schauberger, G. & Tutz, G. (2015). Prediction of Major International Soccer Tournaments Based on '
    'Team-Specific Regularized Poisson Regression. Journal of Quantitative Analysis in Sports, 11(2), 97-115.',
    'Metropolis, N. & Ulam, S. (1949). The Monte Carlo Method. Journal of the American Statistical Association, '
    '44(247), 335-341.',
    'Russell, S. & Norvig, P. (2021). Artificial Intelligence: A Modern Approach (4th ed.). Pearson.',
    'Worldcupwiki.com (2026). FIFA World Cup 2026 - Groups, Schedule and Venues. Consultado: mayo 2026.',
]
for ref in biblio:
    p_normal(ref)

# ============================
# 9. ANEXOS
# ============================
p_h1('Anexos')

p_h2('Anexo A - Dashboard Interactivo (Deploy Vercel)')
p_normal('Aplicacion web con todos los resultados, graficos interactivos y pipeline de agentes:')
p_normal('https://mundial2026-frontend.vercel.app', bold=True)
p_normal(
    'Secciones disponibles desde el menu lateral: Home (/), Favoritos (/favoritos), '
    'Partidos (/partidos), Bracket (/bracket), Comparativa (/comparativa), '
    'ELO Ajustes (/elo-ajustes), Matrices (/matrices), Agentes (/agentes), '
    'Metodologia (/metodologia), Fuentes (/fuentes).'
)

p_h2('Anexo B - Tabla de Convergencia Monte Carlo')
p_normal(
    'N=1,000: Argentina 12.20%, Spain 9.54%, England 9.68%, France 9.04%, Portugal 8.98% | '
    'N=10,000: Argentina 12.45%, Spain 9.31%, England 9.19%, variacion maxima 0.49% | '
    'N=50,000: Argentina 12.56%, Spain 9.33%, England 9.22%, variacion maxima 0.38% (CRITERIO SATISFECHO: < 0.5%).'
)

p_h2('Anexo C - ELOs Ajustados (Top 10 cambios)')
elos = [
    'Morocco: 2435 -> 2490 (+55) | CAF subestimado, #8 FIFA, semifinalista Qatar 2022, campeon AFCON 2026',
    'New Zealand: 1650 -> 1595 (-55) | 18 partidos OFC contra rivales irrelevantes, win rate ficticio',
    'Senegal: 2375 -> 2415 (+40) | Campeon AFCON, sesgo CAF sistematico',
    'Ivory Coast: 2175 -> 2210 (+35) | Diferencial CAF/UEFA injustificado para su nivel',
    'Algeria: 2235 -> 2265 (+30) | Sesgo CAF, rendimiento objetivo superior al ELO base',
    'Japan: 2335 -> 2360 (+25) | Elimino a Alemania y Espana en Qatar 2022, mejor record AFC',
    'Norway: 2250 -> 2275 (+25) | 10 victorias consecutivas, Haaland factor, ranking FIFA subestimado',
    'Belgium: 2470 -> 2445 (-25) | Generacion dorada terminada, De Bruyne 35 anos',
    'Croatia: 2450 -> 2425 (-25) | Modric 40 anos, relegacion Nations League',
    'England: 2520 -> 2500 (-20) | Historial traumatico penales (1990, 1996, 1998, 2004, 2006, Euro 2020)',
]
for e in elos:
    bullet(e, size=11)

p_h2('Anexo D - Archivos de Datos del Proyecto')
p_normal(
    'mundial2026/data/: grupos.json, fixture.json, rankings_fifa.json, elos_equipos.json, elos_ajustados.json, '
    'probabilidades_partidos.json, probabilidades_partidos_v2.json, historial/partidos_oficiales_2018_2026.json, '
    'historial/resumen_por_equipo.json, historial/penales_2018_2026.json.',
    size=11
)
p_normal(
    'mundial2026/outputs/: simulation_results.json (v1 50K torneos), simulation_results_v2.json (v2 FINAL), '
    'resultados_torneo.json (bracket predicho), scores_predichos_v2.json (72 marcadores Poisson).',
    size=11
)
p_normal(
    'Scripts Python: agente_07_scores.py, calc_probabilidades.py, monte_carlo_simulation_v2.py. '
    'Frontend: mundial2026-frontend/src/data/data.ts (fuente de datos React), '
    'mundial2026-frontend/src/pages/ (10 paginas).',
    size=11
)

doc.save(out_path)
print('OK - Guardado en:', out_path)
